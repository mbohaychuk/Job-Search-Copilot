"""Resume text extraction from PDF and DOCX files."""

import hashlib
import io

from jsc.utils.text import normalize_text


class ResumeParser:
    """Extracts plain text from resume documents."""

    def extract_text(self, file_bytes: bytes, content_type: str) -> str:
        """Extract plain text from a PDF or DOCX file.

        Args:
            file_bytes: Raw file bytes.
            content_type: MIME type of the file.

        Returns:
            Cleaned plain text content.
        """
        if "pdf" in content_type:
            return self._extract_pdf(file_bytes)
        elif "wordprocessingml" in content_type or "docx" in content_type:
            return self._extract_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported content type: {content_type}")

    def file_hash(self, file_bytes: bytes) -> str:
        """SHA-256 hash of file bytes."""
        return hashlib.sha256(file_bytes).hexdigest()

    def _extract_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF using PyMuPDF."""
        import pymupdf

        text_parts: list[str] = []
        with pymupdf.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                text_parts.append(page.get_text())
        return normalize_text("\n".join(text_parts))

    def _extract_docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        return normalize_text("\n".join(paragraphs))
